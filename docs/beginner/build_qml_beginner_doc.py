#!/usr/bin/env python3
"""Build the Dutch QML beginner guide as a polished DOCX."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
GUIDE_DIR = Path(__file__).resolve().parent
ASSET_DIR = GUIDE_DIR / "assets"
OUTPUT_DIR = GUIDE_DIR
OUTPUT_PATH = GUIDE_DIR / "qml-van-umi-naar-circuit.docx"
DATA_PATH = ASSET_DIR / "pbmc68k_q4_explanation.json"
CIRCUIT_PATH = ASSET_DIR / "pbmc68k_q4_first_cell_circuit.png"
UNITARY_PATH = ASSET_DIR / "pbmc68k_q4_first_cell_unitary_magnitude.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"
GOLD = "7A5A00"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, *, color="C7CED8", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    total = int(sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_run_font(run, *, name="Calibri", size=None, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_begin, instr_text, fld_char_end))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend((color, underline))
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend((run_pr, text_element))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_caption(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)
    set_keep_with_next(paragraph, False)


def add_equation(doc, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Equation Block")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Cambria Math", size=10.5, color=INK)


def add_code_block(doc, code: str) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    for index, line in enumerate(code.rstrip().splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=8.6, color=INK)


def add_callout(doc, label: str, text: str, *, fill=CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360])
    set_table_borders(table, color=LIGHT_BLUE, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, bold=True, color=INK)
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(text))
            set_run_font(run, color=BLACK)
    set_table_geometry(table, widths)
    return table


def add_custom_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, ind, spacing))
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_element))
    p_pr.append(num_pr)
    paragraph.add_run(text)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        paragraph.paragraph_format.page_break_before = True
    set_keep_with_next(paragraph)


def add_body(doc, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True, color=INK)
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_source_excerpt(
    doc,
    *,
    title: str,
    quote: str,
    meaning: str,
    reference_no: int,
) -> None:
    heading = doc.add_paragraph(style="Heading 2")
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    run = heading.add_run(title)
    set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    set_keep_with_next(heading)
    paragraph = doc.add_paragraph(style="Source Quote")
    run = paragraph.add_run(f'“{quote}”')
    set_run_font(run, name="Calibri", size=10.5, color=INK, italic=True)
    meaning_paragraph = doc.add_paragraph(style="Normal")
    label = meaning_paragraph.add_run("Betekenis voor deze gids: ")
    set_run_font(label, bold=True, color=DARK_BLUE)
    meaning_paragraph.add_run(f"{meaning} [{reference_no}]")


def add_reference(doc, number: int, citation: str, url: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    marker = paragraph.add_run(f"[{number}] ")
    set_run_font(marker, bold=True, color=DARK_BLUE)
    paragraph.add_run(citation + " ")
    add_hyperlink(paragraph, "Online bron", url)


def make_pipeline_figure(path: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    labels = [
        "Biologische\nvraag",
        "UMI-\ncountmatrix",
        "4 waarden\nper cel",
        "RY + CNOT\ncircuit",
        "Z- en ZZ-\nfeatures",
        "Klassieke\nclassifier",
    ]
    figure, axis = plt.subplots(figsize=(11.5, 2.5))
    axis.set_xlim(-0.4, 11.4)
    axis.set_ylim(-0.3, 2.2)
    axis.axis("off")
    xs = np.linspace(0.6, 10.4, len(labels))
    for index, (x, label) in enumerate(zip(xs, labels)):
        color = "#E8EEF5" if index not in (3, 4) else "#D8EAF7"
        box = plt.Rectangle((x - 0.68, 0.58), 1.36, 0.9, facecolor=color, edgecolor="#2E74B5", linewidth=1.7)
        axis.add_patch(box)
        axis.text(x, 1.03, label, ha="center", va="center", fontsize=10, color="#0B2545")
        if index < len(labels) - 1:
            axis.annotate(
                "",
                xy=(xs[index + 1] - 0.75, 1.03),
                xytext=(x + 0.75, 1.03),
                arrowprops={"arrowstyle": "-|>", "color": "#1F4D78", "lw": 1.6},
            )
    axis.text(5.5, 0.08, "Voor iedere cel worden dezelfde stappen herhaald met andere rotatiehoeken.", ha="center", color="#666666", fontsize=9.5)
    figure.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(figure)


def make_umi_figure(path: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    figure, axis = plt.subplots(figsize=(11.5, 4.2))
    axis.set_xlim(0, 12)
    axis.set_ylim(0, 4.4)
    axis.axis("off")

    axis.text(0.4, 3.95, "Oorspronkelijke RNA-moleculen", fontsize=11, color="#0B2545", weight="bold")
    molecules = [(1.2, 3.25, "ACTG1", "UMI ACGT"), (1.2, 1.75, "ACTG1", "UMI TGCA")]
    for x, y, gene, umi in molecules:
        axis.add_patch(plt.Rectangle((x, y), 1.55, 0.62, facecolor="#E8EEF5", edgecolor="#2E74B5", linewidth=1.5))
        axis.text(x + 0.78, y + 0.39, gene, ha="center", va="center", fontsize=9.5, color="#0B2545")
        axis.text(x + 0.78, y + 0.14, umi, ha="center", va="center", fontsize=8.5, color="#666666")

    axis.annotate("PCR-kopiëren", xy=(4.05, 3.4), xytext=(3.0, 3.4), arrowprops={"arrowstyle": "-|>", "color": "#1F4D78"}, color="#666666", ha="center", va="bottom", fontsize=9)
    axis.annotate("PCR-kopiëren", xy=(4.05, 1.9), xytext=(3.0, 1.9), arrowprops={"arrowstyle": "-|>", "color": "#1F4D78"}, color="#666666", ha="center", va="bottom", fontsize=9)

    for row, (y, umi, count) in enumerate(((3.3, "ACGT", 6), (1.8, "TGCA", 3))):
        for copy in range(count):
            x = 4.15 + copy * 0.62
            axis.add_patch(plt.Rectangle((x, y), 0.5, 0.48, facecolor="#F2F4F7", edgecolor="#2E74B5", linewidth=1.0))
            axis.text(x + 0.25, y + 0.24, umi, ha="center", va="center", fontsize=6.7, color="#0B2545")

    axis.annotate("Groepeer gelijke UMI's", xy=(8.55, 2.65), xytext=(7.9, 2.65), arrowprops={"arrowstyle": "-|>", "color": "#1F4D78"}, color="#666666", rotation=90, ha="center", va="center", fontsize=9)
    axis.add_patch(plt.Rectangle((9.15, 2.9), 1.55, 0.62, facecolor="#D8EAF7", edgecolor="#1F4D78", linewidth=1.7))
    axis.add_patch(plt.Rectangle((9.15, 1.75), 1.55, 0.62, facecolor="#D8EAF7", edgecolor="#1F4D78", linewidth=1.7))
    axis.text(9.93, 3.21, "UMI ACGT", ha="center", va="center", fontsize=9.5, color="#0B2545")
    axis.text(9.93, 2.06, "UMI TGCA", ha="center", va="center", fontsize=9.5, color="#0B2545")
    axis.text(9.93, 1.0, "9 reads  →  2 unieke UMI's  →  UMI-telling = 2", ha="center", va="center", fontsize=11, color="#1F4D78", weight="bold")
    figure.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(figure)


def make_feature_figure(path: Path, exact: list[float], sampled: list[float]) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np

    labels = ["Z0", "Z1", "Z2", "Z3", "Z0Z1", "Z1Z2", "Z2Z3", "Z3Z0"]
    x = np.arange(len(labels))
    width = 0.36
    figure, axis = plt.subplots(figsize=(10.7, 4.2))
    axis.axhline(0, color="#C7CED8", linewidth=0.8)
    axis.bar(x - width / 2, exact, width, label="Ideaal", color="#2E74B5")
    axis.bar(x + width / 2, sampled, width, label="512 shots", color="#A9BFD2")
    axis.set_ylim(-1.05, 1.05)
    axis.set_ylabel("verwachtingswaarde")
    axis.set_xticks(x, labels)
    axis.grid(axis="y", color="#E5E7EB", linewidth=0.7)
    axis.spines[["top", "right"]].set_visible(False)
    axis.legend(frameon=False, ncol=2, loc="lower center", bbox_to_anchor=(0.5, 1.01))
    figure.tight_layout()
    figure.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(figure)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    equation = doc.styles.add_style("Equation Block", 1)
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    equation._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    equation.font.size = Pt(10.5)
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.keep_together = True

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.6)
    code.font.color.rgb = rgb(INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    p_pr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)

    source_quote = doc.styles.add_style("Source Quote", 1)
    source_quote.font.name = "Calibri"
    source_quote._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    source_quote._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    source_quote.font.size = Pt(10.5)
    source_quote.font.italic = True
    source_quote.font.color.rgb = rgb(INK)
    source_quote.paragraph_format.left_indent = Inches(0.28)
    source_quote.paragraph_format.right_indent = Inches(0.18)
    source_quote.paragraph_format.space_before = Pt(4)
    source_quote.paragraph_format.space_after = Pt(6)
    source_quote.paragraph_format.line_spacing = 1.20
    source_p_pr = source_quote._element.get_or_add_pPr()
    source_shd = OxmlElement("w:shd")
    source_shd.set(qn("w:fill"), LIGHT_GRAY)
    source_p_pr.append(source_shd)
    source_borders = OxmlElement("w:pBdr")
    source_left = OxmlElement("w:left")
    source_left.set(qn("w:val"), "single")
    source_left.set(qn("w:sz"), "14")
    source_left.set(qn("w:space"), "8")
    source_left.set(qn("w:color"), BLUE)
    source_borders.append(source_left)
    source_p_pr.append(source_borders)


def configure_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "QML BEGINNERSGIDS  |  PBMC68k"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        prefix = paragraph.add_run("Pagina ")
        set_run_font(prefix, size=8.5, color=MUTED)
        add_page_number(paragraph)
        for run in paragraph.runs[1:]:
            set_run_font(run, size=8.5, color=MUTED)


def build_document(data: dict) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    pipeline_path = ASSET_DIR / "qml_pipeline.png"
    umi_path = ASSET_DIR / "umi_counting.png"
    features_path = ASSET_DIR / "quantum_features.png"
    for required in (pipeline_path, umi_path, features_path):
        if not required.exists():
            raise FileNotFoundError(required)

    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "Van UMI-telling naar quantumcircuit"
    doc.core_properties.subject = "Beginnersgids voor een 4-qubit QML-featuremap met PBMC68k-data"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "QML, UMI, PBMC68k, Qiskit, quantumcircuit"
    number_id = add_custom_numbering(doc, bullet=False)
    bullet_id = add_custom_numbering(doc, bullet=True)

    # Cover - editorial_cover pattern with restrained running furniture.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(76)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("EDUCATIEVE QML-GIDS")
    set_run_font(run, size=10, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Van UMI-telling naar quantumcircuit")
    set_run_font(run, size=28, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Hoe je van een biologische classificatievraag naar een 4-qubit QML-circuit gaat")
    set_run_font(run, size=14, color=DARK_BLUE)

    picture = doc.add_picture(str(pipeline_path), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_alt_text(picture, "Stappen van biologische vraag en UMI-countmatrix naar quantumcircuit, quantumfeatures en klassieke classifier")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(22)
    run = meta.add_run("Echte PBMC68k-data | seed 11 | lokale 4-qubit simulator | 512 shots per cel")
    set_run_font(run, size=10.5, color=MUTED, italic=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Versie 21 juli 2026")
    set_run_font(run, size=9.5, color=MUTED)
    add_heading(doc, "1. De hoofdgedachte", 1)
    add_callout(
        doc,
        "Kort gezegd",
        "Een quantumcomputer ontvangt niet automatisch een complete dataset. Voor iedere cel maken we eerst een kleine numerieke vector. Die getallen worden rotatiehoeken in een circuit. Metingen leveren quantumfeatures op, waarna een klassieke classifier de celklasse voorspelt.",
    )
    add_body(
        doc,
        "Dit voorbeeld gebruikt een echte single-cell RNA-countmatrix, maar een lokaal gesimuleerd quantumcircuit. Het is bedoeld om de volledige QML-keten inzichtelijk te maken, niet om quantum advantage aan te tonen.",
    )

    add_heading(doc, "1.1 De biologische ML-vraag", 2)
    add_body(
        doc,
        "De taak is binaire classificatie: bepaal of een geselecteerde bloedcel behoort tot CD4+/CD25 T Reg of CD4+/CD45RO+ Memory. Iedere cel heeft een label y in {-1, +1} en oorspronkelijk 32.738 genkolommen.",
    )
    add_equation(doc, ["cel xᵢ  →  yᵢ ∈ {-1, +1}"])

    add_heading(doc, "1.2 Wat staat er in de PBMC68k-matrix?", 2)
    add_body(
        doc,
        "De volledige matrix bevat 68.579 geannoteerde cellen en 32.738 genen. Een rij is één cel, een kolom is één gen en een matrixelement Xᵢⱼ is de ruwe UMI-telling voor gen j in cel i. De bijbehorende 68k PBMC-profilering is beschreven door Zheng en collega's. [2]",
    )
    add_equation(doc, ["X ∈ ℕ₀^(68.579 × 32.738)", "rij = cel     kolom = gen     Xᵢⱼ = UMI-telling"])
    add_table(
        doc,
        ["Onderdeel", "Betekenis"],
        [
            ["Rij", "Eén afzonderlijke bloedcel"],
            ["Kolom", "Eén gen"],
            ["Getal", "Aantal unieke opgevangen RNA-moleculen"],
            ["Nul", "Geen molecuul gedetecteerd; niet noodzakelijk biologisch afwezig"],
        ],
        [2700, 6660],
    )
    add_caption(doc, "Tabel 1. Hoe de ruwe single-cell countmatrix gelezen wordt.")
    add_heading(doc, "2. UMI-telling: van sequencing reads naar moleculen", 1)
    add_body(
        doc,
        "UMI staat voor Unique Molecular Identifier. Voor het kopiëren krijgt ieder opgevangen RNA-molecuul een kort willekeurig label. PCR kan daarna veel kopieën van hetzelfde molecuul maken, maar al die reads houden dezelfde UMI. Dit principe is ontwikkeld om moleculen vóór amplificatie van elkaar te onderscheiden. [1]",
    )
    picture = doc.add_picture(str(umi_path), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_alt_text(picture, "Diagram waarin negen sequencing reads door gelijke UMI-labels worden teruggebracht tot twee unieke RNA-moleculen")
    add_caption(doc, "Figuur 1. Deduplicatie: negen reads met twee verschillende UMI-labels tellen als twee oorspronkelijke moleculen.")

    add_heading(doc, "2.1 Waarom niet gewoon reads tellen?", 2)
    add_body(
        doc,
        "Zonder UMI's zou een RNA-molecuul dat vaak door PCR is gekopieerd zwaarder meetellen dan een molecuul met weinig kopieën. UMI-deduplicatie probeert deze technische versterking te verwijderen.",
    )
    add_equation(doc, ["UMI-telling ≠ aantal sequencing reads", "UMI-telling ≈ aantal unieke opgevangen RNA-moleculen"])

    add_heading(doc, "2.2 De echte voorbeeldcel", 2)
    cell = data["first_real_training_cell"]
    raw = cell["raw_umi"]
    add_body(
        doc,
        f"De voorbeeldcel is rij {cell['pair_row']} binnen het gefilterde CD4-paar en heeft label {cell['label']}. Over alle 32.738 genen samen bevat deze cel {cell['library_umi']:,} UMI's.".replace(",", "."),
    )
    add_table(
        doc,
        ["Gen", "Ruwe UMI", "Interpretatie"],
        [
            ["IER2", str(raw["IER2"]), "Niet gedetecteerd in deze cel"],
            ["ACTG1", str(raw["ACTG1"]), "Drie unieke opgevangen transcripten"],
            ["LIMD2", str(raw["LIMD2"]), "Eén uniek transcript"],
            ["GLTSCR2", str(raw["GLTSCR2"]), "Twee unieke transcripten"],
        ],
        [1800, 1500, 6060],
    )
    add_caption(doc, "Tabel 2. De vier genwaarden die uiteindelijk naar het 4-qubit-circuit gaan.")
    add_heading(doc, "3. Van 32.738 genen naar vier invoerwaarden", 1)
    add_heading(doc, "3.1 Eerst splitsen, daarna pas voorbereiden", 2)
    add_body(
        doc,
        "De train/test-split moet vóór genselectie en schaling worden gemaakt. In deze demo zijn de vier genen zonder labels gekozen uit alleen de zestien trainingscellen. Ook gemiddelde en standaardafwijking zijn alleen op de trainingscellen bepaald.",
    )
    for item in (
        "Maak een gebalanceerde en disjuncte train/test-split.",
        "Selecteer genen met uitsluitend de trainingsmatrix.",
        "Leer alle schaalparameters uitsluitend op trainingscellen.",
        "Bevries de transformatie en pas haar daarna op de testcellen toe.",
    ):
        add_list_item(doc, item, number_id)

    add_heading(doc, "3.2 Waarom maar vier genen?", 2)
    add_body(
        doc,
        "Bij angle encoding gebruiken we één getal per qubit. Vier qubits kunnen daarom rechtstreeks vier gelijktijdige genwaarden dragen. De geselecteerde genen zijn IER2, ACTG1, LIMD2 en GLTSCR2. Het circuit ziet geen complete cel en geen gennaam: het ontvangt uitsluitend vier numerieke hoeken. De biologische betekenis is dus al sterk klassiek samengevat.",
    )

    add_heading(doc, "3.3 Normalisatie en hoekencoding", 2)
    add_body(
        doc,
        "Cellen kunnen verschillende totale aantallen UMI's hebben. Daarom normaliseren we eerst naar 10.000 counts per cel en nemen we log(1+x). Daarna volgt een trainings-z-score en een begrensde omzetting naar radialen.",
    )
    add_equation(
        doc,
        [
            "x'ᵢⱼ = log(1 + 10⁴ xᵢⱼ / Σₖ xᵢₖ)",
            "zᵢⱼ = (x'ᵢⱼ - μⱼ) / σⱼ",
            "θᵢⱼ = π · clip(zᵢⱼ, -3, 3) / 3",
        ],
    )
    normalized = cell["normalized_log1p"]
    zscore = cell["z_score"]
    angles = cell["rotation_radians"]
    rows = []
    for gene in ("IER2", "ACTG1", "LIMD2", "GLTSCR2"):
        rows.append(
            [
                gene,
                str(raw[gene]),
                f"{normalized[gene]:.4f}",
                f"{zscore[gene]:.4f}",
                f"{angles[gene]:.4f}",
            ]
        )
    add_table(
        doc,
        ["Gen", "UMI", "log1p", "z-score", "θ (rad)"],
        rows,
        [1900, 1200, 1900, 2100, 2260],
    )
    add_caption(doc, "Tabel 3. De volledige omzetting van ruwe telling naar rotatiehoek voor voorbeeldcel 645.")
    add_heading(doc, "4. Van vier hoeken naar een quantumcircuit", 1)
    add_heading(doc, "4.1 Starttoestand", 2)
    add_body(
        doc,
        "Vier qubits starten in |0000>. Omdat vier qubits 2⁴ = 16 basistoestanden hebben, is de toestand een vector met zestien amplitudes. Aan het begin is alleen de eerste amplitude één.",
    )
    add_equation(doc, ["|ψ₀> = |0000> = (1, 0, 0, ..., 0)ᵀ"])

    add_heading(doc, "4.2 RY-poorten schrijven de data in amplitudes", 2)
    add_body(
        doc,
        "Elke geschaalde genwaarde wordt één RY-rotatie. De poort zet |0> om in een superpositie van |0> en |1>. De hoek bepaalt de twee amplitudes. De gebruikte matrix volgt de officiële Qiskit-definitie van RYGate. [3]",
    )
    add_equation(
        doc,
        [
            "Rᵧ(θ) =",
            "⎡ cos(θ/2)   −sin(θ/2) ⎤",
            "⎣ sin(θ/2)    cos(θ/2) ⎦",
            "Rᵧ(θ)|0> = cos(θ/2)|0> + sin(θ/2)|1>",
            "Rᵧ(−1,2187) ≈",
            "⎡  0,8200    0,5723 ⎤",
            "⎣ −0,5723    0,8200 ⎦",
        ],
    )

    add_heading(doc, "4.3 CNOT-poorten maken gezamenlijke features", 2)
    add_body(
        doc,
        "Losse RY-poorten behandelen de vier invoerwaarden onafhankelijk. De CNOT-ring verbindt 0→1, 1→2, 2→3 en 3→0. Hierdoor worden amplitudes en latere meetuitkomsten van meerdere qubits gezamenlijk afhankelijk van de invoerwaarden. In de Qiskit-basis |q₁q₀> heeft CX met q₀ als control en q₁ als target de volgende matrix. [4]",
    )
    add_equation(
        doc,
        [
            "CX₀→₁ =",
            "⎡ 1  0  0  0 ⎤",
            "⎢ 0  0  0  1 ⎥",
            "⎢ 0  0  1  0 ⎥",
            "⎣ 0  1  0  0 ⎦",
            "U(θ) = CX₃→₀ CX₂→₃ CX₁→₂ CX₀→₁ [Rᵧ(θ₃) ⊗ Rᵧ(θ₂) ⊗ Rᵧ(θ₁) ⊗ Rᵧ(θ₀)]",
        ],
    )

    picture = doc.add_picture(str(CIRCUIT_PATH), width=Inches(6.35))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_alt_text(picture, "Vier-qubit Qiskit-circuit met vier RY-poorten, vier CNOT-poorten in een ring en vier metingen")
    add_caption(doc, "Figuur 2. Het daadwerkelijke Qiskit-circuit voor voorbeeldcel 645. De getallen in de RY-poorten zijn de vier berekende hoeken.")
    add_heading(doc, "5. Wat doet de 16 × 16 circuitmatrix?", 1)
    add_body(
        doc,
        "De vier RY-poorten en vier CNOT-poorten vormen samen één unitaire matrix U. Voor vier qubits heeft U zestien rijen en zestien kolommen. Een kolom beschrijft wat het circuit doet met één mogelijke inputbasisstaat.",
    )
    add_equation(doc, ["|ψ(x)> = U(θ)|0000>", "basisvolgorde in Qiskit: |q₃ q₂ q₁ q₀>"])
    add_body(
        doc,
        "Omdat de input |0000> is, is de uiteindelijke toestandsvector precies de eerste kolom van U. De matrix in dit voorbeeld is reëel en numeriek unitair: ||U†U - I||F = 1,1 × 10⁻¹⁵.",
    )
    picture = doc.add_picture(str(UNITARY_PATH), width=Inches(3.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_alt_text(picture, "Heatmap van de absolute waarden van de zestien bij zestien circuitmatrix")
    add_caption(doc, "Figuur 3. Absolute waarden |Uᵢⱼ| van de volledige 16 × 16 matrix. Lichtere vakken hebben een grotere amplitude.")
    add_heading(doc, "5.1 Grootste uitgangskansen", 2)
    probs = cell["largest_output_probabilities"]
    add_table(
        doc,
        ["Bitstring", "Kans", "Betekenis"],
        [
            [probs[0]["basis"], f"{100 * probs[0]['probability']:.2f}%", "Dominante uitgang"],
            [probs[1]["basis"], f"{100 * probs[1]['probability']:.2f}%", "Tweede dominante uitgang"],
            [probs[2]["basis"], f"{100 * probs[2]['probability']:.2f}%", "Kleinere bijdrage"],
        ],
        [1800, 1800, 5760],
    )
    add_caption(doc, "Tabel 4. De grootste ideale meetkansen voor de voorbeeldcel.")
    add_heading(doc, "6. Van metingen naar quantumfeatures", 1)
    add_heading(doc, "6.1 Shots en bitstrings", 2)
    add_body(
        doc,
        "Eén meting levert één bitstring op. Het circuit wordt daarom 512 keer uitgevoerd. Uit de aantallen nullen, enen, gelijke paren en verschillende paren schatten we verwachtingswaarden.",
    )
    add_equation(
        doc,
        [
            "⟨Z_q⟩ = [N(q=0) − N(q=1)] / S",
            "⟨Z_q Z_r⟩ = [N(q=r) − N(q≠r)] / S",
            "S = 512 shots",
        ],
    )
    add_callout(
        doc,
        "Qiskit-bitvolgorde",
        "Qiskit toont een vier-bit-uitkomst als q3 q2 q1 q0. Het meest rechtse bit hoort dus bij qubit 0. Dit is de little-endian conventie die IBM in de Qiskit-documentatie beschrijft. [4]",
        fill=LIGHT_GRAY,
    )

    add_heading(doc, "6.2 De acht gebruikte features", 2)
    add_equation(doc, ["f(x) = (⟨Z₀⟩, ⟨Z₁⟩, ⟨Z₂⟩, ⟨Z₃⟩, ⟨Z₀Z₁⟩, ⟨Z₁Z₂⟩, ⟨Z₂Z₃⟩, ⟨Z₃Z₀⟩)"])
    picture = doc.add_picture(str(features_path), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_alt_text(picture, "Staafdiagram van acht ideale quantumfeatures en de waarden geschat met 512 shots")
    add_caption(doc, "Figuur 4. Ideale verwachtingswaarden tegenover de schatting uit 512 gesimuleerde shots.")

    add_heading(doc, "6.3 Waarom dit circuit klassiek eenvoudig blijft", 2)
    add_body(
        doc,
        "Voor dit specifieke ondiepe RY+CNOT-circuit kunnen de ideale features rechtstreeks als producten van cosinussen worden geschreven. Definieer c_q = cos(θ_q). Dan geldt:",
    )
    add_equation(
        doc,
        [
            "f(x) = (c₁c₂c₃, c₀c₁, c₀c₁c₂, c₀c₁c₂c₃, c₀c₂c₃, c₂, c₃, c₀)",
        ],
    )
    add_body(
        doc,
        "Dit maakt het circuit pedagogisch sterk: we zien precies hoe inputwaarden tot correlaties worden gecombineerd. Tegelijk bewijst het dat deze kleine featuremap geen kandidaat voor quantum advantage is.",
    )
    add_heading(doc, "7. De klassieke classifier en het resultaat", 1)
    add_heading(doc, "7.1 Waarom is dit nog steeds QML?", 2)
    add_body(
        doc,
        "Het quantumcircuit fungeert als featuremap. Voor iedere cel produceert het acht nieuwe getallen. Een klassieke logistische regressie leert vervolgens met de trainingslabels hoe deze features in een voorspelling worden omgezet. Het gebruik van de quantumtoestandsruimte als feature space is een centrale route in quantum machine learning. [5]",
    )
    add_equation(doc, ["P(y=+1 | f) = σ(wᵀf + b)", "σ(a) = 1 / (1 + e^(-a))"])
    add_body(
        doc,
        "De parameters w en b worden klassiek geleerd. Het quantumcircuit zelf is in deze demo niet variabel getraind; de RY-hoeken komen direct uit de celdata en de CNOT-structuur ligt vast.",
    )

    add_heading(doc, "7.2 Gemeten prestatie", 2)
    add_table(
        doc,
        ["Model", "Correct", "Balanced accuracy", "Interpretatie"],
        [
            ["4-qubit quantumfeatures", "7/16", "43,75%", "Onder de klassieke referentie"],
            ["Klassiek, dezelfde vier genen", "9/16", "56,25%", "Beste van deze mini-test"],
        ],
        [2900, 1300, 1900, 3260],
    )
    add_caption(doc, "Tabel 5. Resultaat op de ongeziene testset van zestien cellen.")
    add_callout(
        doc,
        "Conclusie",
        "De keten van echte RNA-data naar quantumfeatures werkt en is volledig reproduceerbaar. Deze kleine demonstratie toont echter geen quantumvoordeel. Zij laat vooral zien wat er bij iedere stap werkelijk met de data gebeurt.",
    )

    add_heading(doc, "7.3 Eén cel is één parametergebonden circuit", 2)
    add_body(
        doc,
        "De 68.579 cellen worden niet tegelijkertijd in vier qubits geladen. Iedere gekozen cel krijgt hetzelfde circuitsjabloon met eigen hoeken. De huidige demonstratie gebruikt zestien trainingscellen en zestien testcellen, dus 32 circuits van elk 512 shots.",
    )
    add_equation(doc, ["cel i  →  θᵢ  →  U(θᵢ)  →  f(xᵢ)"])
    add_heading(doc, "8. Recept: van een nieuw probleem naar een circuit", 1)
    recipe_number_id = add_custom_numbering(doc, bullet=False)
    recipe = (
        "Formuleer precies wat één datapunt, het label en de evaluatiemaatstaf zijn.",
        "Bouw eerst een eerlijke klassieke baseline op dezelfde train/test-splits.",
        "Kies hoeveel klassieke informatie iedere qubit werkelijk moet dragen.",
        "Splits de data voordat je features kiest of schaalparameters leert.",
        "Kies een encoding: angle encoding, amplitude encoding, modules of data re-uploading.",
        "Kies poorten die passen bij de data: enkelvoudige rotaties voor waarden en tweequbitpoorten voor interacties.",
        "Kies vooraf welke observabelen worden gemeten, bijvoorbeeld Z, ZZ of aanvullende X/Y-correlaties.",
        "Test de ideale simulatie, unitariteit, bitvolgorde en shotconvergentie.",
        "Bevries circuit, classifier en hyperparameters voordat je de testset of hardware gebruikt.",
        "Ga pas naar hardware wanneer de ideale route een zinvolle prestatie- of feasibility-vraag beantwoordt.",
    )
    for item in recipe:
        add_list_item(doc, item, recipe_number_id)

    add_heading(doc, "8.1 Compact Qiskit-sjabloon", 2)
    add_code_block(
        doc,
        """from qiskit import QuantumCircuit
def cell_circuit(theta):
    qc = QuantumCircuit(4, 4)
    for qubit, angle in enumerate(theta):
        qc.ry(angle, qubit)
    for control, target in [(0, 1), (1, 2), (2, 3), (3, 0)]:
        qc.cx(control, target)
    qc.measure(range(4), range(4))
    return qc
theta = [-1.218720, 0.385290, 0.278651, 0.098712]
qc = cell_circuit(theta)""",
    )

    add_heading(doc, "8.2 Wat verandert bij 40 of 60 qubits?", 2)
    for item in (
        "Meer qubits helpen alleen als zij meer onafhankelijke en relevante informatie dragen; zij verhogen tegelijk de compilatie-, meet- en foutgevoeligheid.",
        "Een rijkere representatie kan genmodules, meerdere statistieken per module en gerichte langeafstandsinteracties gebruiken. Klassieke voorbereiding en sterke baselines blijven noodzakelijk voor een geloofwaardige voordeelclaim.",
    ):
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "8.3 Relatie tot Quantum Oracle Sketching", 2)
    add_body(
        doc,
        "Dit PBMC68k-beginnersmodel is geen letterlijke QOS-implementatie. Het gebruikt vier klassiek berekende rotatiehoeken, een vaste CNOT-ring en Z/ZZ-readout. De repository bevat daarnaast een afzonderlijke 4-qubit flat-QOS-pilot die wel de officiële q_state_sketch_flat sampling-kern naar Qiskit port. Ook die pilot implementeert slechts één bouwsteen en niet de volledige QOS/QSVT-classificatieketen.",
    )
    add_table(
        doc,
        ["Route", "Doel", "Claimgrens"],
        [
            ["4q PBMC68k-beginnersmodel", "UMI-data begrijpelijk omzetten naar een featuremap en classifier", "Geen letterlijke QOS; geen advantageclaim"],
            ["4q flat-QOS-hardwarepilot", "Officiële sample-afhankelijke phasesketch op hardware testen", "Eén letterlijke QOS-kern; geen complete classifier"],
            ["60q PBMC68k-pilot", "Brede real-data NISQ-featuremap testen", "QOS-geïnspireerd; alleen een afgebakende lokale timingclaim"],
        ],
        [2600, 3900, 2860],
    )
    add_caption(doc, "Tabel 6. De drie experimentele routes mogen niet als hetzelfde algoritme worden beschreven.")

    add_heading(doc, "9. Wat de primaire bronnen zeggen", 1)
    add_body(
        doc,
        "De volgende korte fragmenten zijn in het oorspronkelijke Engels weergegeven. Onder ieder fragment staat waarom die bron relevant is voor deze demonstratie.",
    )
    add_source_excerpt(
        doc,
        title="9.1 Waarom UMI's worden gebruikt",
        quote="unique molecular identifiers (UMIs), which make each molecule in the sample distinct.",
        meaning="Kivioja en collega's leggen het fundamentele idee vast: vóór PCR krijgt ieder molecuul een onderscheidbaar label, zodat kopieën later kunnen worden gededupliceerd.",
        reference_no=1,
    )
    add_source_excerpt(
        doc,
        title="9.2 Waar de PBMC68k-data vandaan komt",
        quote="We profiled 68k peripheral blood mononuclear cells to demonstrate the system's ability to characterize large immune populations.",
        meaning="Zheng en collega's gebruikten deze grote PBMC-verzameling om immuunpopulaties op single-cell-niveau te karakteriseren. Onze twee CD4-klassen komen uit deze route.",
        reference_no=2,
    )
    add_source_excerpt(
        doc,
        title="9.3 De gebruikte rotatiepoort",
        quote="Single-qubit rotation about the Y axis.",
        meaning="De officiële Qiskit-documentatie definieert RY als een rotatie rond de Y-as. In onze featuremap wordt de genwaarde de rotatiehoek.",
        reference_no=3,
    )
    add_source_excerpt(
        doc,
        title="9.4 Waarom bitvolgorde aandacht vraagt",
        quote="qubit 0 is the topmost qubit, and qubit n-1 the bottommost qubit.",
        meaning="Dit verklaart de circuittekening. Bij weergegeven meetbitstrings hanteert Qiskit tegelijk de little-endian conventie, zodat q0 rechts staat.",
        reference_no=4,
    )
    add_source_excerpt(
        doc,
        title="9.5 Het QML-idee achter een featuremap",
        quote="A key component in both methods is the use of the quantum state space as feature space.",
        meaning="Havlíček en collega's beschrijven precies de hybride gedachte die hier wordt onderwezen: de quantumtoestand levert features, terwijl de uiteindelijke classifier klassiek kan blijven.",
        reference_no=5,
    )
    add_heading(doc, "10. Bronnen en reproduceerbaarheid", 1)
    add_reference(
        doc,
        1,
        "Kivioja, T. et al. Counting absolute numbers of molecules using unique molecular identifiers. Nature Methods 9, 72-74 (2012). DOI: 10.1038/nmeth.1778.",
        "https://doi.org/10.1038/nmeth.1778",
    )
    add_reference(
        doc,
        2,
        "Zheng, G. X. Y. et al. Massively parallel digital transcriptional profiling of single cells. Nature Communications 8, 14049 (2017). DOI: 10.1038/ncomms14049.",
        "https://doi.org/10.1038/ncomms14049",
    )
    add_reference(
        doc,
        3,
        "IBM Quantum Documentation. RYGate: single-qubit rotation about the Y axis and its matrix representation.",
        "https://docs.quantum.ibm.com/api/qiskit/qiskit.circuit.library.RYGate",
    )
    add_reference(
        doc,
        4,
        "IBM Quantum Documentation. Bit-ordering in the Qiskit SDK.",
        "https://docs.quantum.ibm.com/guides/bit-ordering",
    )
    add_reference(
        doc,
        5,
        "Havlíček, V. et al. Supervised learning with quantum-enhanced feature spaces. Nature 567, 209-212 (2019). DOI: 10.1038/s41586-019-0980-2.",
        "https://doi.org/10.1038/s41586-019-0980-2",
    )
    add_reference(
        doc,
        6,
        "Zhao, H. et al. Exponential quantum advantage in processing massive classical data. arXiv:2604.07639 (2026).",
        "https://arxiv.org/abs/2604.07639",
    )

    add_heading(doc, "10.1 Brondata", 2)
    p = doc.add_paragraph(style="Normal")
    add_hyperlink(
        p,
        "10x Genomics PBMC68k countmatrix",
        data["dataset"]["matrix_url"],
    )
    p = doc.add_paragraph(style="Normal")
    add_hyperlink(
        p,
        "PBMC68k celtype-annotaties",
        data["dataset"]["annotation_url"],
    )
    add_body(
        doc,
        "Lokale implementatie: qiskit_qos_pbmc68k_q4_educational.py en qiskit_qos_pbmc68k_q4_explain.py. De gebruikte simulatorroute doet geen IBM- of Fire Opal-aanroep.",
    )

    doc.save(OUTPUT_PATH)


def main() -> int:
    for required in (DATA_PATH, CIRCUIT_PATH, UNITARY_PATH):
        if not required.exists():
            raise FileNotFoundError(required)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    build_document(data)
    print(OUTPUT_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
